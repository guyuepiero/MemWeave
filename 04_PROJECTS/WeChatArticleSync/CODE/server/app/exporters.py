"""微文收纳 · 导出引擎（md / json / txt / html / excel / docx）"""
import json
import shutil
from datetime import datetime
from pathlib import Path

import html2text
from openpyxl import Workbook
from docx import Document

from .config import EXPORT_DIR

_H = html2text.HTML2Text()
_H.body_width = 0
_H.ignore_images = False
_H.ignore_emphasis = False
_H.unicode_snob = True


def _fmt_ts(ts):
    if not ts:
        return ""
    try:
        return datetime.fromtimestamp(int(ts)).strftime("%Y-%m-%d %H:%M")
    except Exception:
        return ""


def _slug(title: str, max_len: int = 60) -> str:
    title = re_sub(title)
    return title[:max_len] or "untitled"


def re_sub(s: str) -> str:
    import re

    s = re.sub(r'[\\/:*?"<>|\r\n\t]+', "_", s).strip(" .")
    return s


def md_from_html(content_html: str) -> str:
    return _H.handle(content_html or "").strip()


def _album(a: dict) -> str:
    return a.get("album_name") or (f"专辑 {a.get('album_id','')[-8:]}" if a.get("album_id") else "")


# ---------- 单格式生成 ----------
def build_md(a: dict, content_md: str = "") -> str:
    # 多主题：优先取 topics 数组，逗号分隔；兼容旧单主题字段
    topics = a.get("topics") or []
    topic_str = ", ".join(t.get("name", "") for t in topics if t.get("name")) if topics else (a.get("topic_name") or "无")
    lines = [
        "---",
        f"title: {a.get('title','')}",
        f"author: {a.get('author','')}",
        f"source: {a.get('source','')}",
        f"topic: {topic_str}",
        f"date: {_fmt_ts(a.get('publish_time')) or ''}",
        "origin: 微信公众号",
        f"album: {_album(a)}",
        f"link: {a.get('link','')}",
        f"appmsgid: {a.get('appmsgid','')}",
        f"biz: {a.get('biz','')}",
        "---",
        "",
        f"# {a.get('title','')}",
        "",
        content_md or a.get("content_md", ""),
    ]
    return "\n".join(lines)


def build_json(a: dict) -> dict:
    return {
        "title": a.get("title", ""),
        "author": a.get("author", ""),
        "publish_time": a.get("publish_time"),
        "publish_date": _fmt_ts(a.get("publish_time")),
        "link": a.get("link", ""),
        "source": a.get("source", ""),
        "album": _album(a),
        "album_id": a.get("album_id", ""),
        "topic": a.get("topic_name") or "",
        "appmsgid": a.get("appmsgid", ""),
        "biz": a.get("biz", ""),
        "content_md": a.get("content_md", ""),
        "attrs": json.loads(a.get("attrs_json") or "{}"),
    }


def build_txt(a: dict) -> str:
    md = a.get("content_md", "")
    txt = md.replace("![", "[图](").replace("](images/", "] 本地图片: images/")
    return (
        f"{a.get('title','')}\n"
        f"作者: {a.get('author','')} | 日期: {_fmt_ts(a.get('publish_time'))} | 来源: {a.get('source','')}\n"
        f"所属专辑: {_album(a) or '-'} | 主题: {a.get('topic_name','') or '-'}\n"
        f"链接: {a.get('link','')}\n"
        + "=" * 60
        + "\n\n"
        + txt
    )


def export_excel(articles: list[dict], out_path: Path) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "文章清单"
    ws.append(["标题", "作者", "发布日期", "来源", "所属专辑", "主题", "文章链接", "appmsgid"])
    for a in articles:
        ws.append(
            [
                a.get("title", ""),
                a.get("author", ""),
                _fmt_ts(a.get("publish_time")),
                a.get("source", ""),
                _album(a),
                a.get("topic_name", ""),
                a.get("link", ""),
                a.get("appmsgid", ""),
            ]
        )
    wb.save(out_path)
    return out_path


def export_docx(articles: list[dict], out_path: Path) -> Path:
    doc = Document()
    for a in articles:
        doc.add_heading(a.get("title", ""), level=1)
        doc.add_paragraph(
            f"作者: {a.get('author','')} | 日期: {_fmt_ts(a.get('publish_time'))} | 来源: {a.get('source','')}"
            f" | 所属专辑: {_album(a) or '-'} | 主题: {a.get('topic_name','')}"
        )
        doc.add_paragraph(a.get("link", ""))
        for para in (a.get("content_md", "") or "").splitlines():
            line = para.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=2)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=3)
            elif line.startswith("![") or line.startswith("[图]"):
                continue
            else:
                doc.add_paragraph(line)
    doc.save(out_path)
    return out_path


def export_html(article: dict, out_dir: Path) -> Path:
    """复制 index.html + images/ 到导出目录，保证 100% 还原。"""
    src = Path(article.get("html_path") or "")
    target = out_dir / "index.html"
    if src.exists():
        shutil.copy2(src, target)
        images_src = Path(article.get("images_dir") or "")
        if images_src.exists():
            shutil.copytree(images_src, out_dir / "images", dirs_exist_ok=True)
    return target


# ---------- 批量导出入口 ----------
def run_export(fmt: str, articles: list[dict], tag: str = "all") -> dict:
    """按格式导出文章列表。返回 {dir, count, files}"""
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = EXPORT_DIR / f"{fmt}_{tag}_{stamp}"
    out_dir.mkdir(parents=True, exist_ok=True)
    files = []

    if fmt == "json":
        payload = [build_json(a) for a in articles]
        f = out_dir / "articles.json"
        f.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        files.append(str(f))
    elif fmt == "excel":
        f = out_dir / "文章清单.xlsx"
        export_excel(articles, f)
        files.append(str(f))
    elif fmt == "docx":
        f = out_dir / "文章合集.docx"
        export_docx(articles, f)
        files.append(str(f))
    elif fmt == "html":
        for a in articles:
            d = out_dir / _slug(a.get("source", "unknown")) / _slug(a.get("title", ""))
            d.mkdir(parents=True, exist_ok=True)
            export_html(a, d)
        files.append(str(out_dir))
    elif fmt in ("md", "txt"):
        for a in articles:
            body = build_md(a) if fmt == "md" else build_txt(a)
            ext = ".md" if fmt == "md" else ".txt"
            d = out_dir / _slug(a.get("source", "unknown"))
            d.mkdir(parents=True, exist_ok=True)
            f = d / f"{_fmt_ts(a.get('publish_time'))[:10] or '0000-00-00'}_{_slug(a.get('title',''))}{ext}"
            f.write_text(body, encoding="utf-8")
            files.append(str(f))
    else:
        raise ValueError(f"不支持的格式: {fmt}")

    return {"dir": str(out_dir), "count": len(articles), "files": files}
